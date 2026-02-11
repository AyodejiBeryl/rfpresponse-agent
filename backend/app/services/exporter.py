from __future__ import annotations

import csv
import io
from typing import List

from docx import Document

from app.schemas.analysis import AnalysisResponse, ComplianceRow


def export_csv(rows: List[ComplianceRow]) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["requirement_id", "status", "evidence", "owner", "notes"])
    for row in rows:
        writer.writerow(
            [row.requirement_id, row.status, row.evidence, row.owner or "", row.notes]
        )
    return output.getvalue()


def export_markdown(analysis: AnalysisResponse) -> str:
    md = ["# Proposal Draft", "", f"> {analysis.disclaimer}", "", "## Metadata"]
    for key, value in analysis.metadata.items():
        md.append(f"- **{key}**: {value}")

    md.extend(["", "## Draft Sections"])
    for key, value in analysis.draft_sections.items():
        title = key.replace("_", " ").title()
        md.extend([f"### {title}", value, ""])

    md.extend(["## Gaps"])
    for gap in analysis.gaps:
        md.append(f"- {gap}")

    return "\n".join(md).strip() + "\n"


def export_docx(analysis: AnalysisResponse) -> bytes:
    document = Document()
    document.add_heading("Proposal Draft", level=1)
    document.add_paragraph(analysis.disclaimer)

    document.add_heading("Metadata", level=2)
    for key, value in analysis.metadata.items():
        document.add_paragraph(f"{key}: {value}")

    document.add_heading("Draft Sections", level=2)
    for key, value in analysis.draft_sections.items():
        document.add_heading(key.replace("_", " ").title(), level=3)
        document.add_paragraph(value)

    document.add_heading("Gaps", level=2)
    for gap in analysis.gaps:
        document.add_paragraph(gap, style="List Bullet")

    data = io.BytesIO()
    document.save(data)
    return data.getvalue()
