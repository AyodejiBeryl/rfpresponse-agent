"""
Document export service for generating Word documents from RFP responses.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_rfp_response_document(
    title: str,
    metadata: dict[str, str],
    draft_sections: dict[str, str],
    requirements: list[dict[str, Any]] | None = None,
    compliance_matrix: list[dict[str, Any]] | None = None,
    company_name: str | None = None,
) -> bytes:
    """
    Generate a Word document from RFP response data.

    Returns the document as bytes that can be sent as a file download.
    """
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True

    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True

    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True

    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)

    if company_name:
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run(f"Prepared by: {company_name}")
        company_run.font.size = Pt(14)

    doc.add_paragraph()  # Spacing

    # Add metadata section
    if metadata:
        doc.add_heading("Document Information", level=1)

        # Create a table for metadata
        metadata_items = [
            ("Solicitation Number", metadata.get("solicitation_number", "N/A")),
            ("Due Date", metadata.get("due_date", "N/A")),
            ("Issuing Organization", metadata.get("issuing_organization", "N/A")),
            ("NAICS Code", metadata.get("naics", "N/A")),
            ("Set-Aside", metadata.get("set_aside", "N/A")),
        ]

        table = doc.add_table(rows=len(metadata_items), cols=2)
        table.style = 'Table Grid'

        for i, (label, value) in enumerate(metadata_items):
            if value and value != "N/A" and value != "Not found":
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)

        doc.add_paragraph()  # Spacing

    # Add page break before main content
    doc.add_page_break()

    # Add Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("[Update this table of contents after finalizing the document]")
    doc.add_paragraph()

    # Section order for RFP responses
    section_order = [
        ("executive_summary", "Executive Summary"),
        ("technical_approach", "Technical Approach"),
        ("management_approach", "Management Approach"),
        ("past_performance", "Past Performance"),
        ("staffing_plan", "Staffing Plan"),
        ("pricing_summary", "Pricing Summary"),
    ]

    # Add draft sections in order
    for section_key, section_title in section_order:
        if section_key in draft_sections and draft_sections[section_key]:
            doc.add_heading(section_title, level=1)

            content = draft_sections[section_key]
            # Split content by paragraphs and add each
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a sub-heading (starts with ## or is all caps)
                    if para_text.strip().startswith('##'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        doc.add_heading(heading_text, level=2)
                    elif para_text.strip().startswith('#'):
                        heading_text = para_text.strip().lstrip('#').strip()
                        doc.add_heading(heading_text, level=2)
                    else:
                        # Handle bullet points
                        lines = para_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line.startswith('- ') or line.startswith('* '):
                                # Add as bullet point
                                p = doc.add_paragraph(line[2:], style='List Bullet')
                            elif line.startswith('• '):
                                p = doc.add_paragraph(line[2:], style='List Bullet')
                            elif line:
                                doc.add_paragraph(line)

            doc.add_paragraph()  # Spacing between sections

    # Add any remaining sections not in the standard order
    for section_key, content in draft_sections.items():
        if section_key not in [s[0] for s in section_order] and content:
            # Convert section_key to title case
            section_title = section_key.replace('_', ' ').title()
            doc.add_heading(section_title, level=1)

            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    lines = para_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        elif line:
                            doc.add_paragraph(line)

    # Add compliance matrix as appendix if provided
    if compliance_matrix and len(compliance_matrix) > 0:
        doc.add_page_break()
        doc.add_heading("Appendix A: Compliance Matrix", level=1)

        # Create compliance table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Requirement ID"
        header_cells[1].text = "Status"
        header_cells[2].text = "Evidence"
        header_cells[3].text = "Notes"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for item in compliance_matrix[:50]:  # Limit to 50 rows for document size
            row = table.add_row()
            row.cells[0].text = str(item.get("requirement_id", ""))
            row.cells[1].text = str(item.get("status", "")).upper()
            row.cells[2].text = str(item.get("evidence", ""))[:200]  # Truncate long evidence
            row.cells[3].text = str(item.get("notes", ""))[:100]

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream.getvalue()
